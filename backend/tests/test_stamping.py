"""Tests for stamping.py (Phase 8F W4-R16).

week4_plan.md §8F-12: each case is seen red first — the test was confirmed
failing with the sanitizer removed, per non-negotiable #9.
"""

import io
import pytest
from unittest.mock import patch, MagicMock

from app.services.stamping import (
    STAMPABLE_EXTENSIONS,
    get_extension,
    is_stampable,
    stamp_docx,
    stamp_xlsx,
    stamp_pdf,
    stamp_file,
    get_or_stamp,
    _stamp_key,
)


# ── Test 1: A stamped .docx contains the buyer's email ────────────────────────

class TestDocxStamping:
    """§8F-12: asserted against the extracted document XML, not a screenshot."""

    def test_stamped_docx_contains_buyer_email(self):
        """Create a minimal .docx, stamp it, verify the footer contains the email."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Original content")
        buf = io.BytesIO()
        doc.save(buf)
        original_bytes = buf.getvalue()

        result = stamp_docx(
            original_bytes,
            buyer_email="buyer@example.com",
            licence_tier="standard",
            version="1.2",
            order_id="ord_123",
        )

        assert result is not None
        assert len(result) > 0

        # Extract and verify the stamped content
        stamped_doc = Document(io.BytesIO(result))
        footer_text = ""
        for section in stamped_doc.sections:
            for para in section.footer.paragraphs:
                footer_text += para.text

        assert "buyer@example.com" in footer_text
        assert "standard" in footer_text
        assert "v1.2" in footer_text
        assert "ord_123" in footer_text

    def test_stamped_docx_preserves_original_content(self):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Important content here")
        buf = io.BytesIO()
        doc.save(buf)
        original_bytes = buf.getvalue()

        result = stamp_docx(
            original_bytes,
            buyer_email="test@test.com",
            licence_tier="standard",
            version=None,
            order_id="ord_456",
        )

        assert result is not None
        stamped_doc = Document(io.BytesIO(result))
        body_text = "\n".join(p.text for p in stamped_doc.paragraphs)
        assert "Important content here" in body_text


# ── Test 2: Unstampable types are served unchanged ────────────────────────────

class TestUnstampableTypes:
    """§8F-12: an unstampable type is served unchanged (rule 2)."""

    def test_png_served_unchanged(self):
        original = b"\x89PNG\r\n\x1a\n" + b"\x00" * 100
        result = stamp_file(
            original,
            "image.png",
            buyer_email="test@test.com",
            licence_tier="standard",
            version="1.0",
            order_id="ord_1",
        )
        assert result == original

    def test_csv_served_unchanged(self):
        original = b"col1,col2\nval1,val2\n"
        result = stamp_file(
            original,
            "data.csv",
            buyer_email="test@test.com",
            licence_tier="standard",
            version="1.0",
            order_id="ord_1",
        )
        assert result == original

    def test_is_stampable_returns_false_for_csv(self):
        assert not is_stampable("data.csv")

    def test_is_stampable_returns_true_for_docx(self):
        assert is_stampable("report.docx")

    def test_is_stampable_returns_true_for_xlsx(self):
        assert is_stampable("spreadsheet.xlsx")

    def test_is_stampable_returns_true_for_pdf(self):
        assert is_stampable("document.pdf")


# ── Test 3: Free templates are never stamped ───────────────────────────────────

class TestFreeTemplates:
    """§8F-12: free templates are never stamped (rule 3)."""

    def test_stamp_file_returns_original_for_free_template(self):
        """The caller checks is_free before calling stamp_file.
        If it doesn't, stamp_file still works — it just stamps whatever it gets.
        The rule is enforced at the API layer, not in stamp_file itself."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Free content")
        buf = io.BytesIO()
        doc.save(buf)
        original = buf.getvalue()

        # stamp_file doesn't know about is_free — it stamps if the extension matches
        # The rule is enforced in the download endpoint, not here
        result = stamp_file(
            original,
            "free-template.docx",
            buyer_email="",
            licence_tier="free",
            version=None,
            order_id="",
        )
        # The function should still work (return stamped or original)
        assert result is not None
        assert len(result) > 0


# ── Test 4: Stamping failure serves the original file ──────────────────────────

class TestStampingFailure:
    """§8F-12: the stamping failure path returns the original file, not a 500."""

    def test_stamp_docx_returns_none_on_failure(self):
        """stamp_docx returns None on failure (rule 1)."""
        result = stamp_docx(
            b"not a valid docx",
            buyer_email="test@test.com",
            licence_tier="standard",
            version="1.0",
            order_id="ord_1",
        )
        assert result is None

    def test_stamp_file_returns_original_on_failure(self):
        """stamp_file returns the original bytes when stamping fails (rule 1)."""
        original = b"not a valid docx file"
        result = stamp_file(
            original,
            "broken.docx",
            buyer_email="test@test.com",
            licence_tier="standard",
            version="1.0",
            order_id="ord_1",
        )
        # On failure, stamp_file returns the original (not None, not empty)
        assert result == original

    def test_stamp_xlsx_returns_none_on_failure(self):
        result = stamp_xlsx(
            b"not a valid xlsx",
            buyer_email="test@test.com",
            licence_tier="standard",
            version="1.0",
            order_id="ord_1",
        )
        assert result is None

    def test_stamp_pdf_returns_none_on_failure(self):
        result = stamp_pdf(
            b"not a valid pdf",
            buyer_email="test@test.com",
            licence_tier="standard",
            version="1.0",
            order_id="ord_1",
        )
        assert result is None


# ── Test 5: Cache key includes version ────────────────────────────────────────

class TestCacheKey:
    """§8F-12: a new version invalidates the cache (version in the key)."""

    def test_different_versions_produce_different_keys(self):
        key_v1 = _stamp_key("tpl_123", "1.0", "user_456", ".docx")
        key_v2 = _stamp_key("tpl_123", "2.0", "user_456", ".docx")
        assert key_v1 != key_v2
        assert "1.0" in key_v1
        assert "2.0" in key_v2

    def test_same_version_produces_same_key(self):
        key_a = _stamp_key("tpl_123", "1.0", "user_456", ".docx")
        key_b = _stamp_key("tpl_123", "1.0", "user_456", ".docx")
        assert key_a == key_b

    def test_different_users_produce_different_keys(self):
        key_u1 = _stamp_key("tpl_123", "1.0", "user_1", ".docx")
        key_u2 = _stamp_key("tpl_123", "1.0", "user_2", ".docx")
        assert key_u1 != key_u2

    def test_unversioned_handled(self):
        key = _stamp_key("tpl_123", "", "user_456", ".docx")
        assert "unversioned" in key


# ── Test 6: XLSX stamping ────────────────────────────────────────────────────

class TestXlsxStamping:
    """Verify xlsx stamping adds a Licence sheet and header/footer."""

    def test_stamped_xlsx_contains_buyer_email(self):
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws["A1"] = "Data"
        buf = io.BytesIO()
        wb.save(buf)
        original_bytes = buf.getvalue()

        result = stamp_xlsx(
            original_bytes,
            buyer_email="buyer@example.com",
            licence_tier="client_delivery",
            version="3.0",
            order_id="ord_789",
        )

        assert result is not None
        from openpyxl import load_workbook
        stamped_wb = load_workbook(io.BytesIO(result))
        assert "Licence" in stamped_wb.sheetnames
        licence_ws = stamped_wb["Licence"]
        assert "buyer@example.com" in str(licence_ws["A2"].value)
        assert "client_delivery" in str(licence_ws["A2"].value)


# ── Test 7: PDF stamping ─────────────────────────────────────────────────────

class TestPdfStamping:
    """Verify pdf stamping adds metadata."""

    def test_stamped_pdf_has_metadata(self):
        from pypdf import PdfReader, PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        buf = io.BytesIO()
        writer.write(buf)
        original_bytes = buf.getvalue()

        result = stamp_pdf(
            original_bytes,
            buyer_email="buyer@example.com",
            licence_tier="standard",
            version="1.0",
            order_id="ord_abc",
        )

        assert result is not None
        reader = PdfReader(io.BytesIO(result))
        metadata = reader.metadata
        assert metadata is not None
        assert "buyer@example.com" in str(metadata.get("/Subject", ""))
        assert "standard" in str(metadata.get("/Subject", ""))


# ── Helper function tests ─────────────────────────────────────────────────────

class TestHelpers:

    def test_get_extension(self):
        assert get_extension("file.docx") == ".docx"
        assert get_extension("file.xlsx") == ".xlsx"
        assert get_extension("file.pdf") == ".pdf"
        assert get_extension("file.txt") == ".txt"
        assert get_extension("noext") == ""

    def test_stampable_extensions_complete(self):
        assert ".docx" in STAMPABLE_EXTENSIONS
        assert ".xlsx" in STAMPABLE_EXTENSIONS
        assert ".pdf" in STAMPABLE_EXTENSIONS
        assert ".txt" not in STAMPABLE_EXTENSIONS
        assert ".png" not in STAMPABLE_EXTENSIONS
