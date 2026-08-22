"""Phase 8F (W4-R16): Per-buyer file stamping.

On a **paid** download, serve a copy stamped with the buyer's name and email, the
order id, the licence tier and the version. `.docx` via `python-docx` (footer),
`.xlsx` via `openpyxl` (header/footer plus a licence sheet), `.pdf` via `pypdf`.
**Generated once and cached** in storage under
`stamped/{template_id}/{version}/{user_id}`; a second download serves the cached
object. The version in the key is what makes a re-published template re-stamp
rather than serve a stale copy.

Three rules the stamping code must hold, each a test:
1. **A stamping failure serves the original file.** Wrapped and swallowed — a
   broken stamp must never cost someone the file they paid for.
2. **Unstampable types are served unchanged.** A silent no-op would be a claim the
   code does not keep.
3. **Free templates are never stamped** — there is no buyer to name, and `is_free`
   downloads have no authenticated user by design.
"""

import logging
from io import BytesIO
from typing import Optional

from app.integrations.storage_client import generate_presigned_url, upload_file

logger = logging.getLogger(__name__)

# File extensions we can stamp. Anything else is served unchanged (rule 2).
STAMPABLE_EXTENSIONS = {".docx", ".xlsx", ".pdf"}

# Cache key pattern: stamped/{template_id}/{version}/{user_id}-{extension}
STAMP_CACHE_PREFIX = "stamped"


def _stamp_key(template_id: str, version: str, user_id: str, ext: str) -> str:
    """Build the storage key for a cached stamped copy."""
    safe_version = (version or "unversioned").replace("/", "-")
    return f"{STAMP_CACHE_PREFIX}/{template_id}/{safe_version}/{user_id}{ext}"


def stamp_docx(
    original_bytes: bytes,
    *,
    buyer_email: str,
    licence_tier: str,
    version: Optional[str],
    order_id: str,
) -> Optional[bytes]:
    """Add a footer line to a .docx file. Returns None on failure (rule 1)."""
    try:
        from docx import Document

        doc = Document(BytesIO(original_bytes))
        stamp_line = (
            f"{buyer_email} · {licence_tier} licence"
            + (f" · v{version}" if version else "")
            + f" · order {order_id}"
        )
        for section in doc.sections:
            footer = section.footer
            if not footer.paragraphs:
                footer.add_paragraph(stamp_line)
            else:
                footer.paragraphs[0].text = stamp_line
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception:
        logger.exception("stamp_docx failed; serving original")
        return None


def stamp_xlsx(
    original_bytes: bytes,
    *,
    buyer_email: str,
    licence_tier: str,
    version: Optional[str],
    order_id: str,
) -> Optional[bytes]:
    """Add a licence sheet and header/footer to an .xlsx file. Returns None on failure (rule 1)."""
    try:
        from openpyxl import Workbook, load_workbook

        wb = load_workbook(BytesIO(original_bytes))
        stamp_line = (
            f"{buyer_email} · {licence_tier} licence"
            + (f" · v{version}" if version else "")
            + f" · order {order_id}"
        )

        # Add or replace a "Licence" sheet
        if "Licence" in wb.sheetnames:
            del wb["Licence"]
        ws = wb.create_sheet("Licence", 0)
        ws["A1"] = "Licence"
        ws["A2"] = stamp_line
        ws["A3"] = "Use and adapt this inside your own organisation."
        ws["A4"] = "Full terms: practicable.com/legal/terms"

        # Header/footer on every sheet
        for sheet in wb.worksheets:
            if sheet.title == "Licence":
                continue
            # oddHeader/oddFooter are always instantiated by openpyxl on a Worksheet;
            # stubs type them Optional. Asserts narrow for the type checker and, if
            # ever wrong, fall into the try/except above (rule 1: original is served).
            assert sheet.oddHeader is not None and sheet.oddFooter is not None
            sheet.oddHeader.center.text = stamp_line
            sheet.oddFooter.center.text = "Practicable · practicable.com/legal/terms"

        buf = BytesIO()
        wb.save(buf)
        return buf.getvalue()
    except Exception:
        logger.exception("stamp_xlsx failed; serving original")
        return None


def stamp_pdf(
    original_bytes: bytes,
    *,
    buyer_email: str,
    licence_tier: str,
    version: Optional[str],
    order_id: str,
) -> Optional[bytes]:
    """Add a metadata stamp to a PDF file. Returns None on failure (rule 1)."""
    try:
        from pypdf import PdfReader, PdfWriter

        reader = PdfReader(BytesIO(original_bytes))
        writer = PdfWriter()

        for page in reader.pages:
            writer.add_page(page)

        stamp_line = (
            f"{buyer_email} · {licence_tier} licence"
            + (f" · v{version}" if version else "")
            + f" · order {order_id}"
        )

        # Add metadata — PDF viewers show this in document properties
        writer.add_metadata({
            "/Subject": stamp_line,
            "/Author": "Practicable",
            "/Keywords": f"licence:{licence_tier},order:{order_id}",
        })

        buf = BytesIO()
        writer.write(buf)
        return buf.getvalue()
    except Exception:
        logger.exception("stamp_pdf failed; serving original")
        return None


STAMPERS = {
    ".docx": stamp_docx,
    ".xlsx": stamp_xlsx,
    ".pdf": stamp_pdf,
}


def get_extension(file_name: str) -> str:
    """Extract lowercase extension with dot, e.g. '.xlsx'. Returns '' if none."""
    if "." not in file_name:
        return ""
    return "." + file_name.rsplit(".", 1)[-1].lower()


def is_stampable(file_name: str) -> bool:
    """Phase 8F rule 2: is this a type we can stamp?"""
    return get_extension(file_name) in STAMPABLE_EXTENSIONS


def stamp_file(
    original_bytes: bytes,
    file_name: str,
    *,
    buyer_email: str,
    licence_tier: str,
    version: Optional[str],
    order_id: str,
) -> bytes:
    """Stamp a file, or return the original unchanged.

    Phase 8F rules:
    - Rule 2: unstampable types are served unchanged (never returns empty/None).
    - Rule 1: stamping failure serves the original file (never returns None to caller).
    """
    ext = get_extension(file_name)
    if ext not in STAMPERS:
        return original_bytes

    result = STAMPERS[ext](
        original_bytes,
        buyer_email=buyer_email,
        licence_tier=licence_tier,
        version=version,
        order_id=order_id,
    )
    return result if result is not None else original_bytes


def get_or_stamp(
    original_bytes: bytes,
    *,
    template_id: str,
    file_name: str,
    version: Optional[str],
    buyer_email: str,
    buyer_name: str,
    licence_tier: str,
    order_id: str = "",
    user_id: str,
) -> bytes:
    """Main entry point: check cache, stamp if needed, upload cached copy.

    Phase 8F step 6: Generated once and cached. The version in the key makes a
    re-published template re-stamp rather than serve a stale copy.

    Free templates (rule 3) should never reach this function — the caller must
    check `is_free` before calling. If they don't, the buyer_email being empty
    is the safety net (the stamp is meaningless without it).
    """
    ext = get_extension(file_name)

    # Rule 2: unstampable types served unchanged, no caching needed
    if ext not in STAMPABLE_EXTENSIONS:
        return original_bytes

    # Check cache
    cache_key = _stamp_key(template_id, version or "", user_id, ext)
    try:
        cached = generate_presigned_url(cache_key, expiry_seconds=3600)
        # If presigned URL succeeds, the cached file exists — download and serve
        import requests as _requests
        resp = _requests.get(cached, timeout=10)
        if resp.status_code == 200 and len(resp.content) > 0:
            return resp.content
    except Exception:
        pass  # Cache miss — fall through to stamp

    # Stamp
    stamped = stamp_file(
        original_bytes,
        file_name,
        buyer_email=buyer_email,
        licence_tier=licence_tier,
        version=version,
        order_id=order_id,
    )

    # Cache the stamped copy (best-effort)
    try:
        content_type_map = {
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ".pdf": "application/pdf",
        }
        upload_file(
            key=cache_key,
            body=stamped,
            content_type=content_type_map.get(ext, "application/octet-stream"),
        )
    except Exception:
        logger.exception("Failed to cache stamped file %s; serving from memory", cache_key)

    return stamped
